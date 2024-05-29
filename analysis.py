import os
import numpy as np
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
import PyPDF2
from tkinter import messagebox
from database import save_history
from tkinter import filedialog

current_results = None
current_overall_results = None
query_file_name = None
analysis_saved = False


def preprocess_text(text):
    return text.lower()


def load_files_from_folder(folder_path):
    files_content = {}
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            full_path = os.path.join(folder_path, filename)
            document = Document(full_path)
            text = "\n".join([para.text for para in document.paragraphs])
            files_content[filename] = preprocess_text(text)
        elif filename.endswith(".pdf"):
            full_path = os.path.join(folder_path, filename)
            with open(full_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = "\n".join([page.extract_text() for page in reader.pages])
            files_content[filename] = preprocess_text(text)
        elif filename.endswith(".txt"):
            full_path = os.path.join(folder_path, filename)
            with open(full_path, "r", encoding="utf-8") as file:
                text = file.read()
            files_content[filename] = preprocess_text(text)
    return files_content


def train_tfidf_model(documents):
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(documents)
    return vectorizer, tfidf_matrix


def calculate_similarity(tfidf_matrix, query_vector):
    similarity_scores = cosine_similarity(query_vector, tfidf_matrix)
    return similarity_scores


def calculate_originality(similarity_scores):
    originality_scores = 1 - similarity_scores
    return originality_scores


def calculate_self_similarity(document_text, vectorizer):
    sentences = document_text.split(".")
    sentence_vectors = vectorizer.transform(sentences)
    self_similarity_matrix = cosine_similarity(sentence_vectors)
    np.fill_diagonal(self_similarity_matrix, 0)
    avg_self_similarity = np.mean(self_similarity_matrix)
    return avg_self_similarity


def run_analysis(query_file_path, folder_path, user):
    global current_results, current_overall_results, query_file_name, analysis_saved

    if query_file_path.endswith(".docx"):
        document = Document(query_file_path)
        query_text = preprocess_text(
            "\n".join([para.text for para in document.paragraphs])
        )
    elif query_file_path.endswith(".pdf"):
        with open(query_file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            query_text = preprocess_text(
                "\n".join([page.extract_text() for page in reader.pages])
            )
    elif query_file_path.endswith(".txt"):
        with open(query_file_path, "r", encoding="utf-8") as file:
            query_text = preprocess_text(file.read())

    query_file_name = os.path.basename(query_file_path)

    files_content = load_files_from_folder(folder_path)
    documents = list(files_content.values())
    file_names = list(files_content.keys())

    vectorizer, tfidf_matrix = train_tfidf_model(documents)
    query_vector = vectorizer.transform([query_text])
    similarity_scores = calculate_similarity(tfidf_matrix, query_vector)
    originality_scores = calculate_originality(similarity_scores)

    avg_originality = np.mean(originality_scores) * 100
    avg_citation = np.mean(similarity_scores) * 100
    self_similarity_score = calculate_self_similarity(query_text, vectorizer) * 100

    current_results = pd.DataFrame(
        {
            "File": file_names,
            "Similarity": (similarity_scores.flatten() * 100).round(2),
            "Originality": ((1 - similarity_scores.flatten()) * 100).round(2),
        }
    )
    current_results = current_results.sort_values(by="Similarity", ascending=False)
    current_results.to_csv("plagiarism_results.csv", index=False)

    current_overall_results = {
        "average_originality": avg_originality,
        "average_citation": avg_citation,
        "self_plagiarism": self_similarity_score,
    }

    # Сохранение общих результатов в CSV файл
    overall_results_df = pd.DataFrame([current_overall_results])
    overall_results_df.to_csv("plagiarism_overall_results.csv", index=False)

    save_history(user, query_file_name, current_overall_results)
    analysis_saved = True
    return current_results, current_overall_results
